from flask import Flask, render_template, request, redirect, url_for, flash, send_file
import sqlite3
from config import Config
from docx import Document
import os


app = Flask(__name__)

# Путь для сохранения сгенерированных документов
CONTRACTS_DIR = 'contracts'
os.makedirs(CONTRACTS_DIR, exist_ok=True)


@app.route('/generate_contract', methods=['GET', 'POST'])
def generate_contract():
    # Получение данных номера и даты договора из GET или POST
    contract_number = request.args.get('contract_number') or request.form.get('contract_number', 'UNKNOWN')
    contract_date = request.args.get('contract_date') or request.form.get('contract_date', 'UNKNOWN')

    # Генерация имени и пути к файлу
    file_name = f"contract_{contract_number}_{contract_date}.docx"
    file_path = os.path.join(CONTRACTS_DIR, file_name)

    print(f"Метод: {request.method}, Путь файла: {file_path}")

    if request.method == 'POST':
        # Получение данных из формы
        data = request.form
        print(f"Данные формы: {data}")

        try:
            # Проверяем наличие шаблона
            template_path = 'templates/Абонентский договор.docx'
            if not os.path.exists(template_path):
                print("Шаблон не найден!")
                return "Шаблон не найден!", 500

            # Открываем шаблон и подготавливаем данные для замены
            doc = Document(template_path)
            replacements = {
                '==CONTRACT_NUMBER==': contract_number,
                '==CONTRACT_DATE==': contract_date,
                '==CLIENT_NAME==': data['client_name'],
                '==CLIENT_CONTACT==': data['client_contact'],
                '==CLIENT_DOCUMENT==': data['client_document'],
                '==CLUB_CARD_NUMBER==': data['club_card_number'],
                '==SERVICES_LIST==': data['services_list'],
                '==PAYMENT_METHOD==': data['payment_method'],
                '==CONTRACT_STATUS==': data['contract_status'],
                '==SERVICE_START_DATE==': data['service_start_date'],
                '==SERVICE_END_DATE==': data['service_end_date'],
                '==CLUB_CARD_TYPE==': data['club_card_type'],
                '==ADVANCE_AMOUNT==': data['advance_amount']
            }

            # Замена текста в шаблоне
            for paragraph in doc.paragraphs:
                for placeholder, value in replacements.items():
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, value)

            # Замена текста в таблицах (если есть)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for placeholder, value in replacements.items():
                            if placeholder in cell.text:
                                cell.text = cell.text.replace(placeholder, value)

            # Создание директории для документов (если не существует)
            os.makedirs(CONTRACTS_DIR, exist_ok=True)
            # Сохранение документа
            doc.save(file_path)
            print(f"Файл сохранён: {file_path}")

        except Exception as e:
            # Логирование ошибок
            print(f"Ошибка генерации документа: {e}")
            return "Ошибка генерации документа", 500

        # Отправка файла для скачивания
        return send_file(file_path, as_attachment=True)

    # Если метод GET и файл уже существует, предлагаем его скачать
    if request.method == 'GET' and os.path.exists(file_path):
        print("Отправка существующего файла")
        return send_file(file_path, as_attachment=True)

    # Если GET-запрос и файл не существует, отображаем форму
    print("Форма отображена")
    return render_template('generate_contract.html', contract_number=contract_number, contract_date=contract_date)




# Установите секретный ключ
app.secret_key = 'qwerty12345'
app.config.from_object(Config)

# Функция для взаимодействия с базой данных
def query_db(query, args=(), one=False):
    conn = sqlite3.connect('fitness_club.db')
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.execute(query, args)
    rv = cur.fetchall()
    conn.close()
    return (rv[0] if rv else None) if one else rv

def execute_db(query, args=()):
    conn = sqlite3.connect('fitness_club.db')
    cur = conn.cursor()
    cur.execute(query, args)
    conn.commit()
    conn.close()




# Главная страница
@app.route('/')
def home():
    return render_template('base.html')

# Страница "Клиенты"
@app.route('/clients')
def clients_list():
    clients = query_db('SELECT * FROM clients')
    return render_template('clients_list.html', clients=clients)

# Добавление нового клиента
@app.route('/clients/add', methods=['GET', 'POST'])
def add_client():
    if request.method == 'POST':
        first_name = request.form.get('first_name')
        last_name = request.form.get('last_name')
        date_of_birth = request.form.get('date_of_birth')
        phone_number = request.form.get('phone_number')
        email = request.form.get('email')
        membership_type = request.form.get('membership_type')
        registration_date = request.form.get('registration_date')

        # Минимальная проверка корректности данных
        if not (first_name and last_name and date_of_birth and phone_number and email and membership_type and registration_date):
            flash('Все поля обязательны для заполнения!', 'danger')
            return redirect(url_for('add_client'))

        # Сохранение данных в базу
        execute_db('''
            INSERT INTO clients (first_name, last_name, date_of_birth, phone_number, email, membership_type, registration_date)
            VALUES (?, ?, ?, ?, ?, ?, ?)''',
            (first_name, last_name, date_of_birth, phone_number, email, membership_type, registration_date))

        flash('Клиент успешно добавлен!', 'success')
        return redirect(url_for('clients_list'))

    return render_template('add_client.html')

@app.route('/clients/<int:client_id>')
def client_card(client_id):
    client = query_db('SELECT * FROM clients WHERE client_id = ?', [client_id], one=True)
    return render_template('client_card.html', client=client)

# Страница "Администраторы"
@app.route('/administrators')
def administrators_list():
    administrators = query_db('SELECT * FROM administrators')
    return render_template('administrators_list.html', administrators=administrators)

@app.route('/administrators/add', methods=['GET', 'POST'])
def add_administrator():
    if request.method == 'POST':
        first_name = request.form.get('first_name')
        last_name = request.form.get('last_name')
        hire_date = request.form.get('hire_date')
        phone_number = request.form.get('phone_number')
        email = request.form.get('email')
        shift_schedule = request.form.get('shift_schedule')

        # Минимальная проверка корректности данных
        if not (first_name and last_name and hire_date and phone_number and email and shift_schedule):
            flash('Все поля обязательны для заполнения!', 'danger')
            return redirect(url_for('add_administrator'))

        # Сохранение данных в базу
        execute_db('''
            INSERT INTO administrators (first_name, last_name, hire_date, phone_number, email, shift_schedule)
            VALUES (?, ?, ?, ?, ?, ?)''',
            (first_name, last_name, hire_date, phone_number, email, shift_schedule))

        flash('Администратор успешно добавлен!', 'success')
        return redirect(url_for('administrators_list'))

    return render_template('add_administrator.html')

@app.route('/administrators/<int:admin_id>')
def administrator_card(admin_id):
    administrator = query_db('SELECT * FROM administrators WHERE admin_id = ?', [admin_id], one=True)
    return render_template('administrator_card.html', administrator=administrator)

# Страница "Менеджеры"
@app.route('/managers')
def managers_list():
    managers = query_db('SELECT * FROM managers')
    return render_template('managers_list.html', managers=managers)

@app.route('/managers/add', methods=['GET', 'POST'])
def add_manager():
    if request.method == 'POST':
        first_name = request.form.get('first_name')
        last_name = request.form.get('last_name')
        position = request.form.get('position')
        phone_number = request.form.get('phone_number')
        email = request.form.get('email')

        # Минимальная проверка корректности данных
        if not (first_name and last_name and position and phone_number and email):
            flash('Все поля обязательны для заполнения!', 'danger')
            return redirect(url_for('add_manager'))

        # Сохранение данных в базу
        execute_db('''
            INSERT INTO managers (first_name, last_name, position, phone_number, email)
            VALUES (?, ?, ?, ?, ?)''',
            (first_name, last_name, position, phone_number, email))

        flash('Менеджер успешно добавлен!', 'success')
        return redirect(url_for('managers_list'))

    return render_template('add_manager.html')

@app.route('/managers/<int:manager_id>')
def manager_card(manager_id):
    manager = query_db('SELECT * FROM managers WHERE manager_id = ?', [manager_id], one=True)
    return render_template('manager_card.html', manager=manager)

# Страница "Документы"
@app.route('/documents')
def documents_list():
    documents = query_db('SELECT * FROM documents')
    return render_template('documents_list.html', documents=documents)

@app.route('/documents/add', methods=['GET', 'POST'])
def add_document():
    if request.method == 'POST':
        document_type = request.form.get('document_type')
        creation_date = request.form.get('creation_date')
        related_entity = request.form.get('related_entity')
        status = request.form.get('status')
        file_path = request.form.get('file_path')

        # Минимальная проверка корректности данных
        if not (document_type and creation_date and related_entity and status):
            flash('Все поля обязательны для заполнения!', 'danger')
            return redirect(url_for('add_document'))

        # Сохранение данных в базу
        execute_db('''
            INSERT INTO documents (document_type, creation_date, related_entity, status, file_path)
            VALUES (?, ?, ?, ?, ?)''',
            (document_type, creation_date, related_entity, status, file_path))

        flash('Документ успешно добавлен!', 'success')
        return redirect(url_for('documents_list'))

    return render_template('add_document.html')

@app.route('/documents/<int:document_id>')
def document_card(document_id):
    document = query_db('SELECT * FROM documents WHERE document_id = ?', [document_id], one=True)
    return render_template('document_card.html', document=document)

# Страница "Платежи"
@app.route('/payments')
def payments_list():
    payments = query_db('SELECT * FROM payments')
    return render_template('payments_list.html', payments=payments)

@app.route('/payments/<int:payment_id>')
def payment_card(payment_id):
    payment = query_db('SELECT * FROM payments WHERE payment_id = ?', [payment_id], one=True)
    return render_template('payment_card.html', payment=payment)

@app.route('/payments/add', methods=['GET', 'POST'])
def add_payment():
    if request.method == 'POST':
        client_id = request.form.get('client_id')
        payment_date = request.form.get('payment_date')
        amount = request.form.get('amount')
        payment_method = request.form.get('payment_method')
        status = request.form.get('status')

        # Минимальная проверка корректности данных
        if not (client_id and payment_date and amount and payment_method and status):
            flash('Все поля обязательны для заполнения!', 'danger')
            return redirect(url_for('add_payment'))

        # Сохранение данных в базу
        execute_db('''
            INSERT INTO payments (client_id, payment_date, amount, payment_method, status)
            VALUES (?, ?, ?, ?, ?)''',
            (client_id, payment_date, amount, payment_method, status))

        flash('Платеж успешно добавлен!', 'success')
        return redirect(url_for('payments_list'))

    # Получение списка клиентов для выпадающего списка
    clients = query_db('SELECT client_id, first_name, last_name FROM clients')
    return render_template('add_payment.html', clients=clients)

if __name__ == '__main__':
    app.run(debug=True)
